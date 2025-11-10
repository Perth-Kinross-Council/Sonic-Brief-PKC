import logging
from azure.storage.blob import BlobServiceClient
from azure.identity import ManagedIdentityCredential
from azure.core.exceptions import AzureError
from datetime import datetime

from config import AppConfig

logger = logging.getLogger(__name__)


class StorageService:
    def __init__(self, config: AppConfig):
        self.config = config
        # Use System Assigned Managed Identity for Azure Functions
        # Do not confuse AZURE_CLIENT_ID (Entra ID App Registration) with Managed Identity
        self.credential = ManagedIdentityCredential()

        # Initialize blob service client
        self.blob_service_client = BlobServiceClient(
            account_url=self.config.storage_account_url,
            credential=self.credential,
        )

    def upload_file(self, file_path: str, original_filename: str, case_id: str = None) -> str:
        """Upload a file to blob storage with OWD naming convention"""
        try:
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )

            # Sanitize filename
            sanitized_filename = original_filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
            
            # Generate OWD-compliant blob name
            if case_id:
                # Case-based structure: case_id/audio/filename
                blob_name = f"{case_id}/audio/{sanitized_filename}"
            else:
                # Date-based structure: date/audio/filename  
                current_date = datetime.now().strftime("%Y-%m-%d")
                blob_name = f"{current_date}/audio/{sanitized_filename}"

            blob_client = container_client.get_blob_client(blob_name)

            # Upload the file
            logger.info(f"Uploading file to blob storage: {blob_name}")
            with open(file_path, "rb") as data:
                blob_client.upload_blob(data, overwrite=True)

            return blob_client.url

        except AzureError as e:
            logger.error(f"Azure storage error: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error uploading file: {str(e)}")
            raise

    def upload_text(
        self, container_name: str, blob_name: str, text_content: str
    ) -> str:
        """Upload text content to blob storage"""
        try:
            logger.info(f"Uploading text to: {container_name}/{blob_name}")
            
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            blob_client = container_client.get_blob_client(blob_name)

            blob_client.upload_blob(text_content.encode("utf-8"), overwrite=True)
            
            logger.info(f"Successfully uploaded text to: {blob_client.url}")
            return blob_client.url
        except Exception as e:
            logger.error(f"Error uploading text: {str(e)}")
            raise

    # DEPRECATED: PDF generation method - commented out in favor of DOCX generation
    # Kept for potential future use or rollback scenarios
    # def generate_and_upload_pdf(self, analysis_text: str, blob_name: str) -> str:
    #     """Generate PDF from analysis text and upload to blob storage using blob name"""
    #     try:
    #         logger.info(f"Generating PDF for blob: {blob_name}")
    #         
    #         from reportlab.lib.pagesizes import letter
    #         from reportlab.pdfgen import canvas
    #         import io

    #         # Create PDF in memory
    #         buffer = io.BytesIO()
    #         c = canvas.Canvas(buffer, pagesize=letter)

    #         # Add content to PDF
    #         y = 750  # Starting y position
    #         for line in analysis_text.split("\n"):
    #             if y < 50:  # Start new page if near bottom
    #                 c.showPage()
    #                 y = 750
    #             c.drawString(50, y, line)
    #             y -= 15

    #         c.save()
    #         pdf_content = buffer.getvalue()
    #         
    #         logger.info(f"Generated PDF with {len(pdf_content)} bytes")

    #         # Upload PDF using the blob name provided
    #         container_client = self.blob_service_client.get_container_client(
    #             self.config.storage_recordings_container
    #         )
    #         blob_client = container_client.get_blob_client(blob_name)
    #         
    #         blob_client.upload_blob(pdf_content, overwrite=True)
    #         
    #         logger.info(f"Successfully uploaded PDF to: {blob_client.url}")
    #         return blob_client.url

    #     except Exception as e:
    #         logger.error(f"Error generating/uploading PDF: {str(e)}")
    #         raise

    def generate_and_upload_docx(self, analysis_text: str, blob_name: str) -> str:
        """Generate DOCX from analysis text (markdown/HTML supported) and upload to blob storage using blob name (OWD logic)"""
        try:
            logger.info(f"Generating DOCX for blob: {blob_name}")
            
            import markdown
            from docx import Document
            from bs4 import BeautifulSoup
            import io

            # Convert markdown to HTML
            html_content = markdown.markdown(analysis_text)
            logger.debug(f"Converted markdown to HTML: {len(html_content)} characters")
            
            # Create a new Word Document
            doc = Document()
            
            # Parse HTML and add to DOCX
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for element in soup:
                if element.name == 'h1':
                    doc.add_heading(element.text, level=1)
                elif element.name == 'h2':
                    doc.add_heading(element.text, level=2)
                elif element.name == 'h3':
                    doc.add_heading(element.text, level=3)
                elif element.name == 'p':
                    paragraph = doc.add_paragraph()
                    for child in element.children:
                        if child.name == 'strong':
                            paragraph.add_run(child.text).bold = True
                        elif child.name == 'em':
                            paragraph.add_run(child.text).italic = True
                        else:
                            paragraph.add_run(str(child))  # Convert to string for safety
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.text, style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.text, style='List Number')
                elif hasattr(element, 'text') and element.text.strip():
                    # Handle any remaining text content
                    doc.add_paragraph(element.text)
            
            # Save DOCX to buffer
            finaldocument = io.BytesIO()
            doc.save(finaldocument)
            finaldocument.seek(0)
            docx_content = finaldocument.getvalue()
            
            logger.info(f"Generated DOCX with {len(docx_content)} bytes")
            
            # Upload DOCX using the blob name provided
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_name)
            
            blob_client.upload_blob(docx_content, overwrite=True)
            
            logger.info(f"Successfully uploaded DOCX to: {blob_client.url}")
            return blob_client.url
            
        except Exception as e:
            logger.error(f"Error generating/uploading DOCX: {str(e)}")
            raise
